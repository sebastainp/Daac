from __future__ import print_function
from collections import OrderedDict
import pprint
from docx.oxml.shared import OxmlElement
import docx,socket
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import psutil
import platform
import os
import subprocess
import sys
import re
import time
import getpass
import os.path
import datetime

from datetime import date
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement


from docx import Document
import docx
doc = docx.Document()
doc.add_heading(' General information ')
doc.add_heading(' Purpose ', 1)
doc_para = doc.add_paragraph()


def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def get_size(bytes, suffix="B"):
    """
    Scale bytes to its proper format
    e.g:
        1253656 => '1.20MB'
        1253656678 => '1.17GB'
    """
    factor = 1024
    for unit in ["", "K", "M", "G", "T", "P"]:
        if bytes < factor:
            return f"{bytes:.2f}{unit}{suffix}"
        bytes /= factor

### General Information ####
doc.add_heading(' General information ')
doc.add_heading(' Purpose ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The purpose of this document is to provide detailed technical guidance required to implement a new RedHat Linux Server in accordance with Atos Global Delivery standards and portfolio services. This document describes the requirements and recommendations for the configuration of an RedHat Linux Server managed by Atos.")

doc.add_heading(' Audience ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This document is intended for RedHat Linux Server guidance tasked with implementing or migrating to a new solution. The blueprint assumes that the reader has reasonable grasp of Redhat Enterprise Linux operating system as well as familiarity with architecture principles including high availability and multi-tenancy.")

doc.add_heading(' In Scope', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The scope of Red Hat Enterprise Linux covers the following:")
doc_para = doc.add_paragraph("Supported Operating Systems: RHEL 7.x and  RHEL 8.x", style='List Bullet')
doc_para = doc.add_paragraph("Selecting a technical solution to meet the required SLA", style='List Bullet')
doc_para = doc.add_paragraph("Customizing the infrastructure to meet a solution needs and adhere best practices", style='List Bullet')
doc_para = doc.add_paragraph("RedHat Enterprise Linux deployment", style='List Bullet')
doc_para = doc.add_paragraph("Post-installation configuration and best practices in terms of security and performance", style='List Bullet')
doc_para = doc.add_paragraph("Integration of an RedHat Linux Server with Atos tools (backup, monitoring, reporting, billing)", style='List Bullet')

doc.add_heading(' Out of Scope', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The following areas are explicitly out of scope:")
doc_para = doc.add_paragraph("Unsupported RHEL versions", style='List Bullet')
doc_para = doc.add_paragraph("Application management", style='List Bullet')
doc_para = doc.add_paragraph("Anything not explicitly noted as in scope", style='List Bullet')



doc.add_heading(' Glossary ', 2)
doc_para = doc.add_paragraph()
doc_para.add_run("\n ")
table = doc.add_table(rows=39, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('Abbreviation / Term').bold = True
#shade_cells([heading_cells[0]],"#0066A2")
shade_cells([heading_cells[0]],"#0000FF")
heading_cells[1].paragraphs[0].add_run('Description').bold = True
shade_cells([heading_cells[1]],"#0000FF")

row = table.rows[1]
row.cells[0].text = "BS 7799"
row.cells[1].text = 'Standard for Information Security Management'

row = table.rows[2]
row.cells[0].text = "ITIL"
row.cells[1].text = 'Information Technology Infrastructure Library'

row = table.rows[3]
row.cells[0].text = "LDAP"
row.cells[1].text = 'Lightweight Directory Access Protocol'

row = table.rows[4]
row.cells[0].text = "CMDB"
row.cells[1].text = 'Configuration Management Database'

row = table.rows[5]
row.cells[0].text = "CPU"
row.cells[1].text = 'Central Processing Unit'

row = table.rows[6]
row.cells[0].text = "DR"
row.cells[1].text = 'Disaster Recovery'

row = table.rows[7]
row.cells[0].text = "ICMP"
row.cells[1].text = 'Internet Control Message Protocol'

row = table.rows[8]
row.cells[0].text = "IPA"
row.cells[1].text = 'Integrated Identity and Authentication'

row = table.rows[9]
row.cells[0].text = "HA"
row.cells[1].text = 'High Availability'

row = table.rows[10]
row.cells[0].text = "HLD"
row.cells[1].text = 'High Level Design'

row = table.rows[11]
row.cells[0].text = "LLD"
row.cells[1].text = 'Low Level Design'

row = table.rows[12]
row.cells[0].text = "IDM"
row.cells[1].text = 'Infrastructure Data Management'

row = table.rows[13]
row.cells[0].text = "RPM"
row.cells[1].text = 'RedHat Package Manager'

row = table.rows[14]
row.cells[0].text = "SLA"
row.cells[1].text = 'Service Level Agreement'

row = table.rows[15]
row.cells[0].text = "SMB"
row.cells[1].text = 'Server Message Block'

row = table.rows[16]
row.cells[0].text = "SAN"
row.cells[1].text = 'Storage Area Network'

row = table.rows[17]
row.cells[0].text = "SSL"
row.cells[1].text = 'Secure Sockets Layer'

row = table.rows[18]
row.cells[0].text = "RAID"
row.cells[1].text = 'Redundant Array of Independent Disks'

row = table.rows[19]
row.cells[0].text = "UEFI"
row.cells[1].text = 'Unified Extensible Firmware Interface'

row = table.rows[20]
row.cells[0].text = "EFI"
row.cells[1].text = 'Extensible Firmware Interface'

row = table.rows[21]
row.cells[0].text = "GRUB"
row.cells[1].text = 'GRand Unified Bootloader'

row = table.rows[22]
row.cells[0].text = "NTP"
row.cells[1].text = 'Network Time Protocol'

row = table.rows[23]
row.cells[0].text = "LVM"
row.cells[1].text = 'Logical Volume Manager'

row = table.rows[24]
row.cells[0].text = "SMTP"
row.cells[1].text = 'Simple Mail Transfer Protocol'

row = table.rows[25]
row.cells[0].text = "TCP"
row.cells[1].text = 'Transmission Control Protocol'

row = table.rows[26]
row.cells[0].text = "UDP"
row.cells[1].text = 'User Datagram Protocol'

row = table.rows[27]
row.cells[0].text = "SNMP"
row.cells[1].text = 'Simple Network Management'

row = table.rows[28]
row.cells[0].text = "FCoE"
row.cells[1].text = 'Fiber channel over Ethernet'

row = table.rows[29]
row.cells[0].text = "HBA"
row.cells[1].text = 'Host Bus adaptor'

row = table.rows[30]
row.cells[0].text = "YUM"
row.cells[1].text = 'Yellow Dog Updater, Modified'

row = table.rows[31]
row.cells[0].text = "VLAN"
row.cells[1].text = 'Virtual Local Area Network'

row = table.rows[32]
row.cells[0].text = "VM"
row.cells[1].text = 'Virtual Machine'

row = table.rows[33]
row.cells[0].text = "WAN"
row.cells[1].text = 'Wide Area Network'

row = table.rows[34]
row.cells[0].text = "CMO"
row.cells[1].text = 'Current Mode of Operation'

row = table.rows[35]
row.cells[0].text = "FMO"
row.cells[1].text = 'Future Mode of Operation'

row = table.rows[36]
row.cells[0].text = "SLES"
row.cells[1].text = 'SUSE Linux Enterprise Server'

row = table.rows[37]
row.cells[0].text = "RHEL"
row.cells[1].text = 'Red Hat Enterprise Linux'

row = table.rows[38]
row.cells[0].text = "OL"
row.cells[1].text = 'Oracle Linux'

# add a page break to start a new page
doc.add_page_break()


###
## Server information 
###

doc.add_heading(' Building Blocks ')
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The below table brings up all the basic information about the platforms that sits below the RHEL layer.")
doc_para = doc.add_paragraph()
hostname = socket.gethostname()
serial = subprocess.getoutput('dmidecode -s system-serial-number')
hw = subprocess.getoutput('dmidecode -s system-product-name')
kernel = subprocess.getoutput('uname -r')
uname=platform.uname()
# uptime
# ---------------
with open("/proc/uptime","r") as f:
    uptime=f.read().split(" ")[0].strip()
uptime=int(float(uptime))
uptime_hours=uptime//3600
uptime_minutes=(uptime % 3600) //60

## Table
table = doc.add_table(rows=9, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
#heading_cells[0].text = "Server ID"
#heading_cells[1].text = "Server Details"
heading_cells[0].paragraphs[0].add_run('Server ID').bold = True
shade_cells([heading_cells[0]],"#0000FF")
heading_cells[1].paragraphs[0].add_run('Server Details').bold = True
shade_cells([heading_cells[1]],"#0000FF")

row = table.rows[1]
row.cells[0].text = "System"
row.cells[1].text = (f"{uname.node}")

row = table.rows[2]
row.cells[0].text = "IP Address"
row.cells[1].text =(f"{socket.gethostbyname(hostname)}")

row = table.rows[3]
row.cells[0].text = "OS"
row.cells[1].text = serial
#row.cells[1].text = (f""+str(platform.linux_distribution()))

row = table.rows[4]
row.cells[0].text = "Serial Number"
row.cells[1].text = serial

row = table.rows[5]
row.cells[0].text = "Hardware Type"
row.cells[1].text = hw

row = table.rows[6]
row.cells[0].text = "Kernel Version"
row.cells[1].text = kernel

row = table.rows[7]
row.cells[0].text = "Process"
row.cells[1].text = (f"{uname.processor}")

row = table.rows[8]
row.cells[0].text = "UPTIME"
row.cells[1].text = (f""+str(uptime_hours)+":"+str(uptime_minutes)+"hours")


doc.add_heading(' Installed Software', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about List of Installed Software in the system")
doc_para.add_run("\n ")
package=os.popen("rpm -qa|sort|head -20").read().splitlines()

def convert_los_to_lol(input_list, seprator):
    for idx, row in enumerate(input_list):
        input_list[idx] = row.split(seprator)


def get_row_column_from_lol(input_list):
    return len(input_list), len(input_list[0])


def fill_table(Input_list, Input_table):
    for idxi, rows in enumerate(Input_list):
        if idxi == 0:
            continue
        table_row = Input_table.rows[idxi]
        for idxj, cell in enumerate(rows):
            table_row.cells[idxj].text = cell

convert_los_to_lol(package, ' ')
# Creating Table
#row, col = get_row_column_from_lol(package)
package_table = doc.add_table(50, 5)

#package_table.style = "TableGrid"
package_table.style = "Table Grid"
package_table.allow_autofit = True

#####Customizing Header Section for the packag_table####
hdr_cells = package_table.rows[0].cells

#hdr_cells[0].text = "Packages"
hdr_cells[0].paragraphs[0].add_run('Packages').bold = True
shade_cells([hdr_cells[0]], "#0000FF")

###Fill rest of the package_table with values###
fill_table(package, package_table)
doc_para = doc.add_paragraph()
doc.add_heading(' Running Services', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about List of running Services in the system")
doc_para.add_run("\n")
services=os.popen("systemctl list-units --type service|grep -Ev '^$|LOAD|ACTIVE|SUB|loaded units listed|systemctl list-unit-files|jenkins'|awk '{print $1,$2,$3,$4}'").read().splitlines()

def convert_los_to_lol(input_list, seprator):
    for idx, row in enumerate(input_list):
        input_list[idx] = row.split(seprator)


def get_row_column_from_lol(input_list):
    return len(input_list), len(input_list[0])


def fill_table(Input_list, Input_table):
    for idxi, rows in enumerate(Input_list):
        if idxi == 0:
            continue
        table_row = Input_table.rows[idxi]
        for idxj, cell in enumerate(rows):
            table_row.cells[idxj].text = cell

convert_los_to_lol(services, ' ')

# Creating Table
#row, col = get_row_column_from_lol(services)
services_table = doc.add_table(50, 5)

#services_table.style = "TableGrid"
services_table.style = "Table Grid"
services_table.allow_autofit = True

#####Customizing Header Section for the services_table####
hdr_cells = services_table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('UNIT').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('LOAD').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('ACTIVE').bold = True
shade_cells([hdr_cells[2]], "#0000FF")
hdr_cells[3].paragraphs[0].add_run('SUB').bold = True
shade_cells([hdr_cells[3]], "#0000FF")

###Fill rest of the services_table with values###
fill_table(services, services_table)

doc.add_page_break()

doc.add_heading(' OS Patches ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Below are the list of Security Patches implemented at O/S level. It is recommended to update the hotfixes as soon as it is released ")
doc_para = doc.add_paragraph("Low", style='List Bullet')
doc_para = doc.add_paragraph("Moderate", style='List Bullet')
doc_para = doc.add_paragraph("Important", style='List Bullet')
doc_para = doc.add_paragraph("Critical", style='List Bullet')
doc_para = doc.add_paragraph()

cmd = "dnf updateinfo list --security|egrep -w 'not registered'"
p=subprocess.Popen([cmd], shell=True, stdout=subprocess.PIPE,stderr=subprocess.PIPE)
output = p.communicate()[0]
p1=""+ output.decode().replace('\n','') + ""
with open ("sub.txt","w") as wh:
    wh.write(p1)
wh.close()
with open('sub.txt') as f:
    if 'not registered' in f.read():
        patches=os.popen("dnf updateinfo list --security|grep -Ev  'Low/Sec.|Moderate/Sec.|Updating|metadata'|head -20").read()
        with open("Patch.txt","w") as wh:
            wh.write(patches)
        wh.close()

        patches1=os.popen("cat Patch.txt").read().splitlines()

        ## Create tables
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.allow_autofit = True

        #####Customizing Header Section for the patches_table####
        hdr_cells = table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('Advisories').bold = True
        shade_cells([hdr_cells[0]], "#0000FF")
        hdr_cells[1].paragraphs[0].add_run('Security').bold = True
        shade_cells([hdr_cells[1]], "#0000FF")
        hdr_cells[2].paragraphs[0].add_run('Security-Patches').bold = True
        shade_cells([hdr_cells[2]], "#0000FF")

        ###Fill rest  f the patches_table with values###
        for item in patches1:
            patches2 = item.split()
            row_cells = table.add_row().cells
            row_cells[0].text = patches2[0]
            row_cells[1].text = patches2[1]
            row_cells[2].text = patches2[2]
    else:
        doc_para.add_run("This system is not registered to Red Hat Subscription Management")


doc.add_heading(' Local User Account ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about Local User Account in the system. ")
doc_para = doc.add_paragraph()
users=os.popen("cat /etc/passwd|grep '^root'|awk -F: '{print $1,$3,$4,$6,$7}';awk -F: '$3 >=1000 {print $1,$3,$4,$6,$7}' /etc/passwd|grep -Ev nobody").read().splitlines()

## Create tables
table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.allow_autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('UserName').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('UID').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('GID').bold = True
shade_cells([hdr_cells[2]], "#0000FF")
hdr_cells[3].paragraphs[0].add_run('HomeDirectory').bold = True
shade_cells([hdr_cells[3]], "#0000FF")
hdr_cells[4].paragraphs[0].add_run('DefaultShell').bold = True
shade_cells([hdr_cells[4]], "#0000FF")

for item in users:
    user = item.split()
    row_cells = table.add_row().cells
    row_cells[0].text = user[0]
    row_cells[1].text = user[1]
    row_cells[2].text = user[2]
    row_cells[3].text = user[3]
    row_cells[4].text = user[4]

doc_para = doc.add_paragraph()

doc.add_heading(' Running Processes', 2 )
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about List of running processes in the system. ")
doc_para.add_run("\n")
process=os.popen("ps -eo user,pid,cmd|tail -50|grep -Ev 'tail|ps|awk|sh'|awk '{print $1,$2,$3}'").read().splitlines()

def convert_los_to_lol(input_list, seprator):
    for idx, row in enumerate(input_list):
        input_list[idx] = row.split(seprator)


def get_row_column_from_lol(input_list):
    return len(input_list), len(input_list[0])


def fill_table(Input_list, Input_table):
    for idxi, rows in enumerate(Input_list):
        if idxi == 0:
            continue
        table_row = Input_table.rows[idxi]
        for idxj, cell in enumerate(rows):
            table_row.cells[idxj].text = cell


convert_los_to_lol(process, ' ')


# Creating Table
#row, col = get_row_column_from_lol(process)
process_table = doc.add_table(50, 5)

#process_table.style = "TableGrid"
process_table.style = "Table Grid"
process_table.allow_autofit = True

hdr_cells = process_table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('USER').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('PID').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('CMD').bold = True
shade_cells([hdr_cells[2]], "#0000FF")


###Fill rest of the process_table with values###
fill_table(process, process_table)

doc.add_page_break()
doc.add_heading(' Hardware Orchestration')
doc_para = doc.add_paragraph("This section brings about the behaviour of hardware during the run to have an overview of the operations as a whole. ")
doc.add_heading(' Processors ', 2)
doc_para = doc.add_paragraph()
v_phy=(f" Physical cores: {psutil.cpu_count(logical=False)}")
v_core=(f" Total cores: {psutil.cpu_count(logical=True)}")
cpufreq = psutil.cpu_freq()
v_fre=(f" Current Frequency: {cpufreq.current:.2f}Mhz")

def cpuinfo():
	cpuinfo=OrderedDict()
	procinfo=OrderedDict()

	nprocs = 0
	with open('/proc/cpuinfo') as f:
		for line in f:
			if not line.strip():
                		# end of one processor
				cpuinfo['proc%s' % nprocs] = procinfo
				nprocs=nprocs+1
                		# Reset
				procinfo=OrderedDict()
			else:
				if len(line.split(':')) == 2:
					procinfo[line.split(':')[0].strip()] = line.split(':')[1].strip()
				else:
					procinfo[line.split(':')[0].strip()] = ''
	return cpuinfo

if __name__=='__main__':
	cpuinfo = cpuinfo()
	for processor in cpuinfo.keys():
		cpumodel=(f"Processor: {cpuinfo[processor]['model name']}")
## Table ###

table = doc.add_table(rows=4, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('CPU').bold = True
shade_cells([heading_cells[0]], "#0000FF")
heading_cells[1].paragraphs[0].add_run('CPU Information').bold = True
shade_cells([heading_cells[1]], "#0000FF")

row = table.rows[1]
row.cells[0].text = "CPU Model"
row.cells[1].text = (f"Processor: {cpuinfo[processor]['model name']}")

row = table.rows[2]
row.cells[0].text = "Physical cores"
row.cells[1].text = v_phy

row = table.rows[3]
row.cells[0].text = "Total cores"
row.cells[1].text = v_core

doc_para.add_run("\n")


doc.add_heading(' Storage ', 2)
doc_para = doc.add_paragraph()
# get all disk partitions
partitions = psutil.disk_partitions()
for partition in partitions:
    s1="Device: "+str(partition.device)+"\n"
    s2="Mountpoint: "+str(partition.mountpoint)+"\n"
    s3="File system type: "+str(partition.fstype)+"\n"
    try:
        partition_usage = psutil.disk_usage(partition.mountpoint)
    except PermissionError:
        continue
    s4="Total Size: "+str(get_size(partition_usage.total))+"\n"
    s5="Used: "+str(get_size(partition_usage.used))+"\n"
    s6="Free: "+str(get_size(partition_usage.free))+"\n"
    s7="Percentage: "+str(partition_usage.percent)+"%\n"
    with open("DISK_INFO.txt","a") as wh:
        wh.write(s1)
        wh.write(s2)
        wh.write(s3)
        wh.write(s4)
        wh.write(s5)
        wh.write(s6)
        wh.write(s7)
    wh.close()
	
disk=os.popen("cat DISK_INFO.txt").read().splitlines()

def convert_los_to_lol(input_list, seprator):
    for idx, row in enumerate(input_list):
        input_list[idx] = row.split(seprator)


def get_row_column_from_lol(input_list):
    return len(input_list), len(input_list[0])


def fill_table(Input_list, Input_table):
    for idxi, rows in enumerate(Input_list):
        if idxi == 0:
            continue
        table_row = Input_table.rows[idxi]
        for idxj, cell in enumerate(rows):
            table_row.cells[idxj].text = cell


convert_los_to_lol(disk, ':')


# Creating Table
row, col = get_row_column_from_lol(disk)
disk_table = doc.add_table(row, col)

#disk_table.style = "TableGrid"
disk_table.style = "Table Grid"
disk_table.allow_autofit = True

hdr_cells = disk_table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Name').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('Disk Details').bold = True
shade_cells([hdr_cells[1]], "#0000FF")


###Fill rest of the disk_table with values###
fill_table(disk, disk_table)
os.remove("DISK_INFO.txt")
	
doc.add_heading(' Memory ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The following table shows the current status of memory utilization by Linux system and provides useful analytical information about memory")
doc_para.add_run("\n ")
svmem=psutil.virtual_memory()

### Table ###
table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('Memory').bold = True
shade_cells([heading_cells[0]],"#0000FF")
heading_cells[1].paragraphs[0].add_run('Memory Information').bold = True
shade_cells([heading_cells[1]],"#0000FF")

row = table.rows[1]
row.cells[0].text = "Total Memory"
row.cells[1].text = (f"Total:{get_size(svmem.total)}")

row = table.rows[2]
row.cells[0].text = "Available Memory"
row.cells[1].text = (f"Available:{get_size(svmem.available)}")

row = table.rows[3]
row.cells[0].text = "Used Memory"
row.cells[1].text = (f"Used:{get_size(svmem.used)}")

row = table.rows[4]
row.cells[0].text = "Percentage"
row.cells[1].text = (f"Percentage:{svmem.percent}%")

doc_para = doc.add_paragraph()

doc.add_heading(' SWAP Memory ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The following table shows the current status of virtual memory utilization by Linux system and provides useful analytical information about memory")
doc_para = doc.add_paragraph()
swap=psutil.swap_memory()

### Table ####
table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('SWAP').bold = True
shade_cells([heading_cells[0]],"#0000FF")
heading_cells[1].paragraphs[0].add_run('Swap Memory Information').bold = True
shade_cells([heading_cells[1]],"#0000FF")

row = table.rows[1]
row.cells[0].text = "Total Memory"
row.cells[1].text = (f"Total:{get_size(swap.total)}")

row = table.rows[2]
row.cells[0].text = "Available Memory"
row.cells[1].text = (f"Free:{get_size(swap.free)}")

row = table.rows[3]
row.cells[0].text = "Used Memory"
row.cells[1].text = (f"Used:{get_size(swap.used)}")

row = table.rows[4]
row.cells[0].text = "Percentage"
row.cells[1].text = (f"Percentage:{swap.percent}%")

doc.add_heading(' Network ', 2)
doc_para = doc.add_paragraph()
net_io = psutil.net_io_counters()
if_addrs = psutil.net_if_addrs()
for interface_name, interface_addresses in if_addrs.items():
    for address in interface_addresses:
        if str(address.family) == 'AddressFamily.AF_INET':
            s2="Interface;  "+str(interface_name)+"\n"
            s3="IP Address;  "+str(address.address)+"\n"
            s4="Netmask;  "+str(address.netmask)+"\n"
            s5="Broadcast IP;  "+str(address.broadcast)+"\n"
            with open("NETWORK_INFO.txt","a") as wh:
                #print("Name;  Network Information; Network-IO-Byte-Receive; Network-IO-Byte-Sent", file=wh)
                wh.write(s2)
                wh.write(s3)
                wh.write(s4)
                wh.write(s5)
            wh.close()
        elif str(address.family) == 'AddressFamily.AF_PACKET':
            sx="MAC Address;  "+str(address.address)+"\n"
            with open("NETWORK_INFO.txt","a") as wh:
                wh.write(sx)
            wh.close()
			
network=os.popen("cat NETWORK_INFO.txt").read().splitlines()

def convert_los_to_lol(input_list, seprator):
    for idx, row in enumerate(input_list):
        input_list[idx] = row.split(seprator)


def get_row_column_from_lol(input_list):
    return len(input_list), len(input_list[0])


def fill_table(Input_list, Input_table):
    for idxi, rows in enumerate(Input_list):
        if idxi == 0:
            continue
        table_row = Input_table.rows[idxi]
        for idxj, cell in enumerate(rows):
            table_row.cells[idxj].text = cell


convert_los_to_lol(network, ';')


# Creating Table
row, col = get_row_column_from_lol(network)
network_table = doc.add_table(row, col)

network_table.style = "Table Grid"
network_table.allow_autofit = True

hdr_cells = network_table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Name').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('Network Details').bold = True
shade_cells([hdr_cells[1]], "#0000FF")

###Fill rest of the network_table with values###
fill_table(network, network_table)
os.remove("NETWORK_INFO.txt")

doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph()
doc.add_heading(' Network Route', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about Route in the system. ")
doc_para.add_run("\n ")
#networkr=os.popen("/usr/sbin/route -n|grep -Ev 'Kernel|Destination'").read().splitlines()

## Create tables
table = doc.add_table(rows=1, cols=8)
table.style = "Table Grid"
table.allow_autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Destination').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('Gateway').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('Genmask').bold = True
shade_cells([hdr_cells[2]], "#0000FF")
hdr_cells[3].paragraphs[0].add_run('Flags').bold = True
shade_cells([hdr_cells[3]], "#0000FF")
hdr_cells[4].paragraphs[0].add_run('Metric').bold = True
shade_cells([hdr_cells[4]], "#0000FF")
hdr_cells[5].paragraphs[0].add_run('Ref').bold = True
shade_cells([hdr_cells[5]], "#0000FF")
hdr_cells[6].paragraphs[0].add_run('Use').bold = True
shade_cells([hdr_cells[6]], "#0000FF")
hdr_cells[7].paragraphs[0].add_run('Iface').bold = True
shade_cells([hdr_cells[7]], "#0000FF")



doc.add_heading(' Services and Ports ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The following table is the true information about the ports used by application and it's associated services.")
doc_para.add_run("\n ")
#servicep=os.popen("netstat -tulp|grep -Ev Active|grep LISTEN").read().splitlines()

## Create tables
table = doc.add_table(rows=1, cols=7)
table.style = "Table Grid"
table.allow_autofit = True


hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Proto').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('Recv-Q').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('Send-Q').bold = True
shade_cells([hdr_cells[2]], "#0000FF")
hdr_cells[3].paragraphs[0].add_run('Local-Address').bold = True
shade_cells([hdr_cells[3]], "#0000FF")
hdr_cells[4].paragraphs[0].add_run('Foreign-Address').bold = True
shade_cells([hdr_cells[4]], "#0000FF")
hdr_cells[5].paragraphs[0].add_run('State').bold = True
shade_cells([hdr_cells[5]], "#0000FF")
hdr_cells[6].paragraphs[0].add_run('PID/Program').bold = True
shade_cells([hdr_cells[6]], "#0000FF")



doc.add_page_break()

doc.add_heading(' RHEL')
doc.add_heading(' Current Login Details', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section contains the Current login information in RHEL server ")
doc_para.add_run("\n ")
login=os.popen("who|grep -Ev tty").read().splitlines()

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.allow_autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Login').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('Terminal').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('Login-Date').bold = True
shade_cells([hdr_cells[2]], "#0000FF")
hdr_cells[3].paragraphs[0].add_run('Login-Time').bold = True
shade_cells([hdr_cells[3]], "#0000FF")
hdr_cells[4].paragraphs[0].add_run('Remote-Server').bold = True
shade_cells([hdr_cells[4]], "#0000FF")

for item in login:
    loginc = item.split()
    row_cells = table.add_row().cells
    row_cells[0].text = loginc[0]
    row_cells[1].text = loginc[1]
    row_cells[2].text = loginc[2]
    row_cells[3].text = loginc[3]
    row_cells[4].text = loginc[4]

doc.add_heading(' SELinux', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("SELinux should be enabled in Permissive mode on all systems as per the Atos Global standard. Current status can be found by running the following command:")
doc_para.add_run("\n ")
selinux=os.popen("sestatus").read().splitlines()

## Create tables
table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
table.allow_autofit = True

for item in selinux:
    selinx = item.split(":")
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('SELinux status').bold = True
    row_cells[1].text = selinx[1]

doc.add_heading(' Firewall', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The system should be configured to use Firewalld to allow and deny access.")
doc_para.add_run("\n ")
#firewalls=os.popen("firewall-cmd --state|grep -v '^$'").read().splitlines()

## Create tables
table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
table.allow_autofit = True


doc_para = doc.add_paragraph()

#fd=os.popen("firewall-cmd --zone=public --list-all").read()
table = doc.add_table(rows=2, cols=1)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('Firewalld Configuration:').bold = True
shade_cells([heading_cells[0]], "#0000FF")
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Here services can be enabled to allow incoming connections. The above shows the default of primary interface firewalled and only port 22 open.")

doc.add_page_break()

doc.add_heading(' Target/Runlevel ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The default run level should be set to be multi-user.target")
doc_para = doc.add_paragraph()
doc_para.add_run("\n ")
doc_para.add_run("Verify default runlevel").bold = True
doc_para.add_run("\n ")
rn=os.popen("systemctl get-default").read()
rn1="systemctl set-default multi-user.target"
table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('Default Runlevel').bold = True
shade_cells([heading_cells[0]],"#0000FF")
heading_cells[1].paragraphs[0].add_run('Command To Change The Runlevel').bold = True
shade_cells([heading_cells[1]],"#0000FF")
row = table.rows[1]
row.cells[0].text = (rn)
row.cells[1].text = (rn1)

#doc.add_page_break()

doc.add_heading(' Chrony/NTP', 2)
doc_para = doc.add_paragraph("chrony is an implementation of the Network Time Protocol. It's a replacement for the ntpd.")
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The output displays information about the current time sources that chronyd is accessing.")
doc_para = doc.add_paragraph(".")
doc_para.add_run("\n ")
ch=subprocess.check_output('rpm -qa chrony | awk -F"-" \'{print $1}\'',shell=True)
ch1=""+ ch.decode().replace('\n','') +""
if ch1 == 'chrony':
    ntp=os.popen("chronyc tracking").read().splitlines()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = True

    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Name').bold = True
    shade_cells([hdr_cells[0]], "#0000FF")
    hdr_cells[1].paragraphs[0].add_run('Tracking').bold = True
    shade_cells([hdr_cells[1]], "#0000FF")

    for item in ntp:
        ntp1 = item.split(":")
        row_cells = table.add_row().cells
        row_cells[0].text =  ntp1[0]
        row_cells[1].text =  ntp1[1]
else:
    doc_para.add_run("chrony package is not Installed")

doc.add_heading(' TimeZone', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about current TimeZone in the system. ")
doc_para.add_run("\n ")

tz=os.popen("timedatectl|grep 'Time zone'").read().splitlines()

table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
table.allow_autofit = True

for item in tz:
    tz1 = item.split(":")
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('Time zone').bold = True
    row_cells[1].text = tz1[1]

doc.add_page_break()

doc.add_heading(' Cluster ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("RedHat Cluster are often used to support mission-critical services in the enterprise. The major goal of cluster is to keep services as available as possible by eliminating bottlenecks and single points of failure. ")
doc_para = doc.add_paragraph()
doc_para.add_run("\n ")
doc_para.add_run("Cluster Status:").bold = True
import pathlib
file = pathlib.Path("/etc/corosync/corosync.conf")
if file.exists ():
    cluster=os.popen("pcs cluster status").read().splitlines()
    with open("pcs.txt","w") as wh:
    	wh.write(cluster)
    wh.close()
    cluster1=os.popen("cat pcs.txt|grep -Ev 'Cluster Status'").read().splitlines()
    ## Create tables
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.allow_autofit = True
	
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Cluster Status:').bold = True
    shade_cells([hdr_cells[0]], "#0000FF")
	
    for item in cluster1:
        row_cells = table.add_row().cells
        row_cells[0].text = item
else:
    doc_para.add_run("Cluster is not Configured").bold = True

doc.add_heading(' Log File ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Log files are files that contain messages about the system, including the kernel, services, and applications running on it. There are different log files for different information. For example:")
doc_para = doc.add_paragraph("/var/log/messages: This file has all the global system messages located inside, including the messages that are logged during system startup. Depending on how the syslog config file is sent up, there are several things that are logged in this file including mail, cron, daemon, kern, auth, etc.", style='List Bullet')
doc_para = doc.add_paragraph("/var/log/secure: Contains information related to authentication and authorization privileges. For example, sshd logs all the messages here, including unsuccessful login", style='List Bullet')
#doc_para = doc.add_paragraph("/var/log/boot.log: Contains information that are logged when the system boots", style='List Bullet')
doc_para = doc.add_paragraph()
m=os.popen('ls -lh /var/log/messages|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Live Log File"}\'').read()
s=os.popen('ls -lh /var/log/secure|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Live Log File"}\'').read()
#b=os.popen('ls -lh /var/log/boot.log|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Live Log File"}\'').read()
m1=os.popen('ls -lh /var/log/messages-*|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Archive Log File"}\'').read()
s1=os.popen('ls -lh /var/log/secure-*|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Archive Log File"}\'').read()
#b1=os.popen('ls -lh /var/log/boot.log-*|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Archive Log File"}\'').read()
#with open("LOG.txt","w") as wh:
#    wh.write(m+s+b+m1+s1+b1)
#wh.close()

#log=os.popen("cat LOG.txt").read().splitlines()

## Create tables
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
table.allow_autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Log_File').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('Date').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('Size').bold = True
shade_cells([hdr_cells[2]], "#0000FF")
hdr_cells[3].paragraphs[0].add_run('File Type').bold = True
shade_cells([hdr_cells[3]], "#0000FF")

#for item in log:
#    log1 = item.split(";")
#    row_cells = table.add_row().cells
#    row_cells[0].text = log1[0]
#    row_cells[1].text = log1[1]
#    row_cells[2].text = log1[2]
#    row_cells[3].text = log1[3]

#doc.add_page_break()

doc.add_heading(' Atos Technology Framework ')
doc_para = doc.add_paragraph("The Atos Technology Framework provides a tooling solution for the Atos Service Management Model (ASMM) and the associated processes and consistently manages the interactions between all components, the Services and all users based on a flexible IT architecture. Henkel tooling will be removed, and Atos tooling will be installed.")
doc.add_heading(' Monitoring Tool ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph()

doc.add_heading(' Patching Tool ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Linux servers will be patched using BMC blade logic tool. BMC BladeLogic agent RPM will be installed in the server and GDTS BladeLogic team will add that server into the management. RPM should be received from GDTS team")
doc_para.add_run("\n")
bladelogic=os.popen("rpm -qa |grep BladeLogic").read()
if len(bladelogic) != 0:
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.allow_autofit = True
    # populate header row -------
    heading_cells = table.rows[0].cells
    heading_cells[0].paragraphs[0].add_run('BMC Agent').bold = True
    shade_cells([heading_cells[0]],"#0000FF")
    row = table.rows[1]
    row.cells[0].text = (bladelogic)
    doc_para = doc.add_paragraph()
    doc_para.add_run("\n ")
    doc_para.add_run("BMC Agent Status").bold = True
    doc_para.add_run("\n")
	
    bladelogic1=os.popen("ps aux|grep rscd |grep -v grep|awk '{print $1,$2,$8,$11}'").read().splitlines()

	## Create tables
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.allow_autofit = True
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('USER').bold = True
    shade_cells([hdr_cells[0]], "#0000FF")
    hdr_cells[1].paragraphs[0].add_run('PID').bold = True
    shade_cells([hdr_cells[1]], "#0000FF")
    hdr_cells[2].paragraphs[0].add_run('STAT').bold = True
    shade_cells([hdr_cells[2]], "#0000FF")
    hdr_cells[3].paragraphs[0].add_run('Service').bold = True
    shade_cells([hdr_cells[3]], "#0000FF")
	
    for item in bladelogic1:
        bl = item.split()
        row_cells = table.add_row().cells
        row_cells[0].text = bl[0]
        row_cells[1].text = bl[1]
        row_cells[2].text = bl[2]
        row_cells[3].text = bl[3]
else:
    doc_para.add_run("BladeLogic package is not installed").bold = True

doc.add_heading(' Backup Tool ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("EMC Networker agent will be installed on Linux servers")
doc_para.add_run("\n")
networker=os.popen("rpm -qa |grep -i lgto").read()
if len(networker) != 0:
    nwservice=os.popen("systemctl is-active networker").read()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = True
    # populate header row -------
    heading_cells = table.rows[0].cells
    heading_cells[0].paragraphs[0].add_run('NetWorker Client').bold = True
    shade_cells([heading_cells[0]],"#0000FF")
    heading_cells[1].paragraphs[0].add_run('NetWorker Service').bold = True
    shade_cells([heading_cells[1]],"#0000FF")
    row = table.rows[1]
    row.cells[0].text = (networker)
    row.cells[1].text = (nwservice)
else:
    doc_para.add_run("Networker Agent is not installed").bold = True

doc.add_heading(' System Recovery Tool ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Relax-and-Recover(ReaR) is a recovery and system migration utility. The utility produces a bootable image and restores from backup using this image. It also allows to restore to different hardware and can therefore be used as a migration utility as well.")
doc_para.add_run("\n")
table = doc.add_table(rows=2, cols=1)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('REAR Version').bold = True
shade_cells([heading_cells[0]],"#0000FF")
row = table.rows[1]
row.cells[0].text = ('rear')

doc.add_page_break()

doc.add_heading(' Health Checkup ')
doc_para = doc.add_paragraph("The following data table is very important during hardening and onboarding of RHEL Server instance. It provides Precise information about server health based on Information Security Standards. ")
doc_para.add_run("\n ")
#os.system("python3 Linux_Helath_Check.py")
#ht=os.popen("cat HEALTH.txt").read().splitlines()

## Create tables
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.allow_autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('NAME').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('THRESHOLD(%)').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('STATUS').bold = True
shade_cells([hdr_cells[2]], "#0000FF")

#for item in ht:
#    Green = parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w')))
#    Red = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
#    Orange = parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w')))
#
#    ht1 = item.split()
#    row_cells = table.add_row().cells
#    row_cells[0].text = ht1[0]
#    row_cells[1].text = ht1[1]
#    #row_cells[2].text = ht1[2]
#    if ht1[2] == 'OK':
        #row_cells[2].paragraphs[0].add_run(ht1[2]).font.color.rgb =  RGBColor(0, 255, 0)
#        row_cells[2].paragraphs[0].add_run(ht1[2]).bold = True
#        row_cells[2]._tc.get_or_add_tcPr().append(Green)
#    elif ht1[2] == 'WARNING':
        #row_cells[2].paragraphs[0].add_run(ht1[2]).font.color.rgb =  RGBColor(255, 140, 0)
#        row_cells[2].paragraphs[0].add_run(ht1[2]).bold = True
#        row_cells[2]._tc.get_or_add_tcPr().append(Orange)
#    else:
#        #row_cells[2].paragraphs[0].add_run(ht1[2]).font.color.rgb =  RGBColor(255, 0, 0)
#        row_cells[2].paragraphs[0].add_run(ht1[2]).bold = True
#        row_cells[2]._tc.get_or_add_tcPr().append(Red)

doc.add_page_break()

doc.add_heading(' Atos TSS ')
doc_para = doc.add_paragraph("Atos standard build procedure will be used to harden the system. Every new server installed should be secured following the Unix Security Standards.  Note that there may be customer requirements which mean that configurations listed in the policy may need to be altered.  Any such alterations should be documented as exceptions to the Unix Security Standards in the server documentation.")
doc_para.add_run("\n ")
#os.system("/usr/bin/python3 tss-script-new.py")
#tss1=os.popen("cat TSS.txt").read().splitlines()

## Create tables
table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.allow_autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Measure ID').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('Measure Title').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('TSS Recommendation').bold = True
shade_cells([hdr_cells[2]], "#0000FF")
hdr_cells[3].paragraphs[0].add_run('Current Value').bold = True
shade_cells([hdr_cells[3]], "#0000FF")
hdr_cells[4].paragraphs[0].add_run('Compliance/Non-Compliance').bold = True
shade_cells([hdr_cells[4]], "#0000FF")



doc.add_page_break()

doc.add_heading(' Risk Analysis ')
doc_para = doc.add_paragraph("The following data table is very important during hardening and onboarding of RHEL Servers. It provides Deep Dive Risk information that must be taken into consideration . ")
doc_para.add_run("\n ")
#os.system("/usr/bin/python3 risk-ana-script-new.py")
#riska=os.popen("cat RISK_ANA.txt").read().splitlines()

## Create tables
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
table.allow_autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Risk Measure').bold = True
shade_cells([hdr_cells[0]], "#0000FF")
hdr_cells[1].paragraphs[0].add_run('Risk Assessment').bold = True
shade_cells([hdr_cells[1]], "#0000FF")
hdr_cells[2].paragraphs[0].add_run('Assessment Result').bold = True
shade_cells([hdr_cells[2]], "#0000FF")
hdr_cells[3].paragraphs[0].add_run('Risk Severity').bold = True
shade_cells([hdr_cells[3]], "#0000FF")


doc.add_page_break()

doc.add_heading(' Backup ')
doc_para = doc.add_paragraph("EMC Networker agent will be installed on servers, under '/nsr' to take all drives/partition backup and it will be backed up in accordance with the existing Atos standard regime.")
doc_para = doc.add_paragraph("RMAN backup tool will be used to take Oracle DB backups.", style='List Bullet')
doc_para = doc.add_paragraph("A full backup will be done once a week, incremental backups will be done for the rest of the week. Backup retention period will be one month or as per Atos backup policy standard.", style='List Bullet')
doc_para.add_run("\n ")
mount=os.popen("df -h|grep -v 'File'|awk '{print $6}'|grep -Ev 'run|dev|sys|mnt'").read()
vg=os.popen("vgs").read()
with open("vg.txt","w") as wh:
    wh.write(vg)
    wh.close()
v = open("vg.txt","r")
for line in v:
    fields = line.split()
    f1=fields[5]
    f2 = ""+str(f1)+""
    f3 = ""+str(f2.replace('<', ' ').replace('VSize',' '))+""
v.close()
table = doc.add_table(rows=3, cols=5)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('Service Component').bold = True
shade_cells([heading_cells[0]],"#0000FF")
heading_cells[1].paragraphs[0].add_run('Devices').bold = True
shade_cells([heading_cells[1]],"#0000FF")
heading_cells[2].paragraphs[0].add_run('Data Volumes (GB)').bold = True
shade_cells([heading_cells[2]],"#0000FF")
heading_cells[3].paragraphs[0].add_run('Agent Type').bold = True
shade_cells([heading_cells[3]],"#0000FF")
heading_cells[4].paragraphs[0].add_run('Archive Frequency').bold = True
shade_cells([heading_cells[4]],"#0000FF")
row = table.rows[1]
row.cells[0].text = "Operating system , all drives/partition Local"
row.cells[1].text = (mount)
row.cells[2].text = "f3"
row.cells[3].text = "EMC Networker"
row.cells[4].text = "Daily Incremental\Weekly Full"
row = table.rows[2]
row.cells[0].text = "Operating system , all drives/partition on SAN"
row.cells[1].text = "Not Applicable"
row.cells[2].text = "Not Applicable"
row.cells[3].text = "EMC Networker"
row.cells[4].text = "Daily Incremental\Weekly Full"

doc.add_page_break()

doc.add_heading(' System Recovery ')
doc_para = doc.add_paragraph("System recovery processes are used to restore a server after some failure be it hardware or software, it can be regarded as part of Disaster Recovery (DR) or as a separate entity if a customer does not have DR. There are several methods of system recovery listed below:")
doc_para = doc.add_paragraph()
doc_para.add_run("\n")
doc_para.add_run("Relax and Recover (REAR)").bold = True
doc_para.add_run("\n")
doc_para.add_run("Relax-and-Recover produces a bootable image which can recreate the system’s original storage layout. Once that is done it initiates a restore from backup. Since the storage layout can be modified prior to recovery, and dissimilar hardware and virtualization is supported, Relax-and-Recover offers the flexibility to be used for complex system migrations.")
doc_para = doc.add_paragraph()
doc_para.add_run("\n")
doc_para.add_run("Enterprise system restore (NetBackup / Networker etc) ").bold = True
doc_para.add_run("\n")
doc_para.add_run("An enterprise solution can be used to restore a system such as NetBackup (Symantec) or Networker (EMC).")

doc.add_page_break()

doc.add_heading(' Redhat Licensing ')
doc_para = doc.add_paragraph("Red Hat Subscriptions (or licensing) are available in many forms, the 2 common ones are:")
doc_para = doc.add_paragraph("RED HAT ENTERPRISE LINUX SERVER", style='List Bullet')
doc_para = doc.add_paragraph("RED HAT ENTERPRISE LINUX FOR VIRTUAL DATACENTERS", style='List Bullet')
doc_para = doc.add_paragraph()
doc_para.add_run("\n")
doc_para.add_run("Table 1. Service-level agreements for Red Hat Enterprise Linux subscriptions.").bold = True
doc_para.add_run("\n")
table = doc.add_table(rows=8, cols=4)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('Service').bold = True
shade_cells([heading_cells[0]],"#0000FF")
heading_cells[1].paragraphs[0].add_run('Self-support').bold = True
shade_cells([heading_cells[1]],"#0000FF")
heading_cells[2].paragraphs[0].add_run('Standard').bold = True
shade_cells([heading_cells[2]],"#0000FF")
heading_cells[3].paragraphs[0].add_run('Premium').bold = True
shade_cells([heading_cells[3]],"#0000FF")
row = table.rows[1]
row.cells[0].text = "Hours of coverage"
row.cells[1].text = "None"
row.cells[2].text = "Standard business"
row.cells[3].text = "Standard business hours (24x7 for Severity 1 and Severity 2)"
row = table.rows[2]
row.cells[0].text = "Support channel"
row.cells[1].text = "None"
row.cells[2].text = "Web and phone"
row.cells[3].text = "NA"
row = table.rows[3]
row.cells[0].text = "Number of cases"
row.cells[1].text = "None"
row.cells[2].text = "Unlimited"
row.cells[3].text = "Unlimited"
row = table.rows[4]
row.cells[0].text = "Severity 1 (Urgent)"
row.cells[1].text = "1 business hour"
row.cells[2].text = "1 hour"
row.cells[3].text = "1 hour"
row = table.rows[5]
row.cells[0].text = "Severity 2 (High)"
row.cells[1].text = "4 business hours"
row.cells[2].text = "2 business hours"
row.cells[3].text = "4 hours"
row = table.rows[6]
row.cells[0].text = "Severity 3 (Medium)"
row.cells[1].text = "1 business day"
row.cells[2].text = "4 business hours"
row.cells[3].text = "8 business hours"
row = table.rows[7]
row.cells[0].text = "Severity 4 (Low)"
row.cells[1].text = "2 business days"
row.cells[2].text = "8 business hours"
row.cells[3].text = "2 business days or as agreed"


doc.add_page_break()

doc.add_heading(' Redhat LifeCycle ')
doc.add_heading(' End of Life ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Red Hat offers subscription services for each major release of Red Hat Enterprise Linux throughout four life-cycle phases—called Full Support, Maintenance Support 1, Maintenance Support 2, and an Extended Life Phase.")
doc_para.add_run("\n")
doc_para = doc.add_paragraph(" Please visit RedHat link mentioned below for Red Hat Enterprise Linux Life Cycle policies.")
#doc_para.add_run("\n")
p = doc.add_paragraph()
add_hyperlink(p, 'Red Hat Enterprise Linux Life Cycle', "https://access.redhat.com/support/policy/updates/errata")
doc_para.add_run("\n")
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph(" The content has been taken from offline")
doc_para = doc.add_paragraph()
table = doc.add_table(rows=2, cols=4)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].paragraphs[0].add_run('Products Released').bold = True
shade_cells([heading_cells[0]],"#0000FF")
heading_cells[1].paragraphs[0].add_run('Lifecycle Start Date').bold = True
shade_cells([heading_cells[1]],"#0000FF")
heading_cells[2].paragraphs[0].add_run('Full support ends').bold = True
shade_cells([heading_cells[2]],"#0000FF")
heading_cells[3].paragraphs[0].add_run('Maintenance Support ends').bold = True
shade_cells([heading_cells[3]],"#0000FF")


row = table.rows[1]
row.cells[0].text = "Redhat Enterprise Linux Server 8"
row.cells[1].text = "May 7, 2019"
row.cells[2].text = "May 31, 2024"
row.cells[3].text = "May 31, 2029"

#doc.add_page_break()
e = datetime.datetime.now()
d=e.strftime("%Y-%m-%d_%H-%M-%S")
doc.save("Redhat_LLD_"+str(d)+".docx")

