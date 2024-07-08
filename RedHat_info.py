#!/usr/bin/python3
from __future__ import print_function
from collections import OrderedDict
import pprint
from docx.oxml.shared import OxmlElement
import docx,socket
#print (os.uname())
#import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import psutil
import platform
import os
import subprocess
import sys
import re
import time
import getpass
import os.path
import datetime

from datetime import date
from docx.enum.dml import MSO_THEME_COLOR_INDEX

from docx import Document
doc = docx.Document('template.docx')
#style = doc.styles['Normal']
#style = doc.styles['WD_STYLE.BODY_TEXT']
doc_para = doc.add_paragraph()

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def get_size(bytes, suffix="B"):
    """
    Scale bytes to its proper format
    e.g:
        1253656 => '1.20MB'
        1253656678 => '1.17GB'
    """
    factor = 1024
    for unit in ["", "K", "M", "G", "T", "P"]:
        if bytes < factor:
            return f"{bytes:.2f}{unit}{suffix}"
        bytes /= factor

### General Information ####
doc.add_heading(' General information ')
doc.add_heading(' Purpose ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The purpose of this document is to provide detailed technical guidance required to implement a new RedHat Linux Server in accordance with Atos Global Delivery standards and portfolio services. This document describes the requirements and recommendations for the configuration of an RedHat Linux Server managed by Atos.")

doc.add_heading(' Audience ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This document is intended for RedHat Linux Server guidance tasked with implementing or migrating to a new solution. The blueprint assumes that the reader has reasonable grasp of Redhat Enterprise Linux operating system as well as familiarity with architecture principles including high availability and multi-tenancy.")

doc.add_heading(' In Scope', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The scope of Red Hat Enterprise Linux covers the following:")
doc_para = doc.add_paragraph("Supported Operating Systems: RHEL 6.x, RHEL 7.x and  RHEL 8.x", style='List Bullet')
doc_para = doc.add_paragraph("Selecting a technical solution to meet the required SLA", style='List Bullet')
doc_para = doc.add_paragraph("Customizing the infrastructure to meet a solution needs and adhere best practices", style='List Bullet')
doc_para = doc.add_paragraph("RedHat Enterprise Linux deployment", style='List Bullet')
doc_para = doc.add_paragraph("Post-installation configuration and best practices in terms of security and performance", style='List Bullet')
doc_para = doc.add_paragraph("Integration of an RedHat Linux Server with Atos tools (backup, monitoring, reporting, billing)", style='List Bullet')

doc.add_heading(' Out of Scope', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The following areas are explicitly out of scope:")
doc_para = doc.add_paragraph("Unsupported RHEL versions", style='List Bullet')
doc_para = doc.add_paragraph("Application management", style='List Bullet')
doc_para = doc.add_paragraph("Anything not explicitly noted as in scope", style='List Bullet')


doc.add_heading(' Glossary ', 2)
doc_para = doc.add_paragraph()
doc_para.add_run("\n ")
table = doc.add_table(rows=30, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Abbreviation / Term"
heading_cells[1].text = "Description"

row = table.rows[1]
row.cells[0].text = "BS 7799"
row.cells[1].text = 'Standard for Information Security Management'

row = table.rows[2]
row.cells[0].text = "ITIL"
row.cells[1].text = 'Information Technology Infrastructure Library'

row = table.rows[3]
row.cells[0].text = "LDAP"
row.cells[1].text = 'Lightweight Directory Access Protocol'

row = table.rows[4]
row.cells[0].text = "CMDB"
row.cells[1].text = 'Configuration Management Database'

row = table.rows[5]
row.cells[0].text = "CPU"
row.cells[1].text = 'Central Processing Unit'

row = table.rows[6]
row.cells[0].text = "DR"
row.cells[1].text = 'Disaster Recovery'

row = table.rows[7]
row.cells[0].text = "ICMP"
row.cells[1].text = 'Internet Control Message Protocol'

row = table.rows[8]
row.cells[0].text = "IPA"
row.cells[1].text = 'Integrated Identity and Authentication'

row = table.rows[9]
row.cells[0].text = "HA"
row.cells[1].text = 'High Availability'

row = table.rows[10]
row.cells[0].text = "HLD"
row.cells[1].text = 'High Level Design'

row = table.rows[11]
row.cells[0].text = "LLD"
row.cells[1].text = 'Low Level Design'

row = table.rows[12]
row.cells[0].text = "IDM"
row.cells[1].text = 'Infrastructure Data Management'

row = table.rows[13]
row.cells[0].text = "RPM"
row.cells[1].text = 'RedHat Package Manager'

row = table.rows[14]
row.cells[0].text = "SLA"
row.cells[1].text = 'Service Level Agreement'

row = table.rows[15]
row.cells[0].text = "SMB"
row.cells[1].text = 'Server Message Block'

row = table.rows[16]
row.cells[0].text = "SAN"
row.cells[1].text = 'Storage Area Network'

row = table.rows[17]
row.cells[0].text = "SSL"
row.cells[1].text = 'Secure Sockets Layer'

row = table.rows[18]
row.cells[0].text = "RAID"
row.cells[1].text = 'Redundant Array of Independent Disks'

row = table.rows[19]
row.cells[0].text = "UEFI"
row.cells[1].text = 'Unified Extensible Firmware Interface'

row = table.rows[20]
row.cells[0].text = "EFI"
row.cells[1].text = 'Extensible Firmware Interface'

row = table.rows[21]
row.cells[0].text = "GRUB"
row.cells[1].text = 'GRand Unified Bootloader'

row = table.rows[22]
row.cells[0].text = "NTP"
row.cells[1].text = 'Network Time Protocol'

row = table.rows[23]
row.cells[0].text = "LVM"
row.cells[1].text = 'Logical Volume Manager'

row = table.rows[24]
row.cells[0].text = "SMTP"
row.cells[1].text = 'Simple Mail Transfer Protocol'

row = table.rows[25]
row.cells[0].text = "TCP"
row.cells[1].text = 'Transmission Control Protocol'

row = table.rows[26]
row.cells[0].text = "UDP"
row.cells[1].text = 'User Datagram Protocol'

row = table.rows[27]
row.cells[0].text = "SNMP"
row.cells[1].text = 'Simple Network Management'

row = table.rows[28]
row.cells[0].text = "FCoE"
row.cells[1].text = 'Fiber channel over Ethernet'

row = table.rows[29]
row.cells[0].text = "HBA"
row.cells[1].text = 'Host Bus adaptor'

row = table.rows[30]
row.cells[0].text = "YUM"
row.cells[1].text = 'Yellow Dog Updater, Modified'

row = table.rows[31]
row.cells[0].text = "VLAN"
row.cells[1].text = 'Virtual Local Area Network'

row = table.rows[32]
row.cells[0].text = "VM"
row.cells[1].text = 'Virtual Machine'

row = table.rows[33]
row.cells[0].text = "WAN"
row.cells[1].text = 'Wide Area Network'

row = table.rows[34]
row.cells[0].text = "CMO"
row.cells[1].text = 'Current Mode of Operation'

row = table.rows[35]
row.cells[0].text = "FMO"
row.cells[1].text = 'Future Mode of Operation'

row = table.rows[36]
row.cells[0].text = "SLES"
row.cells[1].text = 'SUSE Linux Enterprise Server'

row = table.rows[37]
row.cells[0].text = "RHEL"
row.cells[1].text = 'Red Hat Enterprise Linux'

row = table.rows[38]
row.cells[0].text = "OL"
row.cells[1].text = 'Oracle Linux'

# add a page break to start a new page
doc.add_page_break()

###
## Server information 
###

doc.add_heading(' Building Blocks ')
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The below table brings up all the basic information about the platforms that sits below the RHEL layer.")
#doc_para.add_run("\n")
hostname = socket.gethostname()
serial = subprocess.getoutput('dmidecode -s system-serial-number')
hw = subprocess.getoutput('dmidecode -s system-product-name')
kernel = subprocess.getoutput('uname -r')
uname=platform.uname()
# uptime
# ---------------
with open("/proc/uptime","r") as f:
    uptime=f.read().split(" ")[0].strip()
uptime=int(float(uptime))
uptime_hours=uptime//3600
uptime_minutes=(uptime % 3600) //60

## Table
table = doc.add_table(rows=9, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Server ID"
heading_cells[1].text = "Server Details"

row = table.rows[1]
row.cells[0].text = "System"
row.cells[1].text = (f"{uname.node}")

row = table.rows[2]
row.cells[0].text = "IP Address"
row.cells[1].text =(f"{socket.gethostbyname(hostname)}")

row = table.rows[3]
row.cells[0].text = "OS"
row.cells[1].text = (f""+str(platform.linux_distribution()))

row = table.rows[4]
row.cells[0].text = "Serial Number"
row.cells[1].text = serial

row = table.rows[5]
row.cells[0].text = "Hardware Type"
row.cells[1].text = hw

row = table.rows[6]
row.cells[0].text = "Kernel Version"
row.cells[1].text = kernel

row = table.rows[7]
row.cells[0].text = "Process"
row.cells[1].text = (f"{uname.processor}")

row = table.rows[8]
row.cells[0].text = "UPTIME"
row.cells[1].text = (f""+str(uptime_hours)+":"+str(uptime_minutes)+"hours")


doc.add_heading(' Installed Software', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about List of Installed Software in the system")
doc_para.add_run("\n ")
package=os.popen("/bin/rpm -qa|sort|head -20").read()
with open("PACKAGE.txt","w") as wh:
    print("Package", file=wh)
    wh.write(package)
wh.close()
f = open("PACKAGE.txt","r")
for line in f:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (line)
f.close()

doc.add_heading(' Running Services', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about List of running Services in the system")
doc_para.add_run("\n")
#r=os.popen("/bin/bash sed_task.sh").read()
#r=int(r)
#if(r<=6):
#    service=os.popen("chkconfig --list").read()
#    doc_para.add_run(service)
#    doc_para.add_run("\n")
#else:
s=os.popen("systemctl list-units --type service|grep -Ev '^$|LOAD|ACTIVE|SUB|loaded units listed|systemctl list-unit-files|jenkins'|awk '{print $1,$2,$3,$4}'").read()
with open("SERVICE.txt","w") as wh:
    print("UNIT  LOAD  ACTIVE  SUB", file=wh)
    wh.write(s)
wh.close()
f = open("SERVICE.txt","r")
for line in f:
    fields = line.strip().split()
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    f4=fields[3]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
    row.cells[3].text = (f4)
f.close()

doc.add_page_break()

doc.add_heading(' OS Patches ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Below are the list of Security Patches implemented at O/S level. It is recommended to update the hotfixes as soon as it is released ")
doc_para = doc.add_paragraph("Low", style='List Bullet')
doc_para = doc.add_paragraph("Moderate", style='List Bullet')
doc_para = doc.add_paragraph("Important", style='List Bullet')
doc_para = doc.add_paragraph("Critical", style='List Bullet')
doc_para = doc.add_paragraph()
#r=os.popen("dnf updateinfo list security installed|grep -Ev  'Low/Sec.|Moderate/Sec.|Updating|metadata'|head -20").read()
#with open("patch.txt","w") as wh:
#    print("Advisories" " " "Security" " " "Security-Patches",file=wh)
#    wh.write(r)
#wh.close()
f = open("patch.txt","r")
for line in f:
    fields = line.split(" ")
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
f.close()

doc.add_heading(' Local User Account ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about Local User Account in the system. ")
doc_para = doc.add_paragraph()
user=os.popen("awk -F: '$3 >=1000 {print $1,$3,$4,$6,$7}' /etc/passwd|grep -Ev nobody").read()
su=os.popen("awk -F: '$3 ==0 {print $1,$3,$4,$6,$7}' /etc/passwd").read()
with open("USER.txt","w") as wh:
    print("UserName  UID  GID  HomeDirectory  DefaultShell", file=wh)
    wh.write(su+user)
wh.close()
f = open("USER.txt","r")
for line in f:
    user = line.split()
    f1=user[0]
    f2=user[1]
    f3=user[2]
    f4=user[3]
    f5=user[4]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
    row.cells[3].text = (f4)
    row.cells[4].text = (f5)
f.close()

doc.add_heading(' Running Processes', 2 )
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about List of running processes in the system. ")
doc_para.add_run("\n")
p=os.popen("/usr/bin/ps -eo user,pid,cmd|tail -50|grep -Ev 'tail|/usr/bin/ps|awk|sh'|awk '{print $1,$2,$3}'").read()
with open("PROCESS.txt","w") as wh:
#    print("USER  PID  CMD", file=wh)
    wh.write(p)
wh.close()
f = open("PROCESS.txt","r")
for line in f:
    fields = line.split(" ")
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
f.close()
doc.add_page_break()
doc.add_heading(' Hardware Orchestration')
doc_para = doc.add_paragraph("This section brings about the behaviour of hardware during the run to have an overview of the operations as a whole. ")
doc.add_heading(' Processors ', 2)
doc_para = doc.add_paragraph()
v_phy=(f" Physical cores: {psutil.cpu_count(logical=False)}")
v_core=(f" Total cores: {psutil.cpu_count(logical=True)}")
cpufreq = psutil.cpu_freq()
v_fre=(f" Current Frequency: {cpufreq.current:.2f}Mhz")

def cpuinfo():
	cpuinfo=OrderedDict()
	procinfo=OrderedDict()

	nprocs = 0
	with open('/proc/cpuinfo') as f:
		for line in f:
			if not line.strip():
                		# end of one processor
				cpuinfo['proc%s' % nprocs] = procinfo
				nprocs=nprocs+1
                		# Reset
				procinfo=OrderedDict()
			else:
				if len(line.split(':')) == 2:
					procinfo[line.split(':')[0].strip()] = line.split(':')[1].strip()
				else:
					procinfo[line.split(':')[0].strip()] = ''
	return cpuinfo

if __name__=='__main__':
	cpuinfo = cpuinfo()
	for processor in cpuinfo.keys():
		cpumodel=(f"Processor: {cpuinfo[processor]['model name']}")
## Table ###

table = doc.add_table(rows=4, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = "CPU"
heading_cells[1].text = "CPU Information"

row = table.rows[1]
row.cells[0].text = "CPU Model"
row.cells[1].text = (f"Processor: {cpuinfo[processor]['model name']}")

row = table.rows[2]
row.cells[0].text = "Physical cores"
row.cells[1].text = v_phy

row = table.rows[3]
row.cells[0].text = "Total cores"
row.cells[1].text = v_core

doc_para.add_run("\n")
#doc_para = doc.add_paragraph(" ")

doc.add_heading(' Storage ', 2)
doc_para = doc.add_paragraph()
# get all disk partitions
partitions = psutil.disk_partitions()
for partition in partitions:
    s1="Device: "+str(partition.device)+"\n"
    s2="Mountpoint: "+str(partition.mountpoint)+"\n"
    s3="File system type: "+str(partition.fstype)+"\n"
    try:
        partition_usage = psutil.disk_usage(partition.mountpoint)
    except PermissionError:
        continue
    s4="Total Size: "+str(get_size(partition_usage.total))+"\n"
    s5="Used: "+str(get_size(partition_usage.used))+"\n"
    s6="Free: "+str(get_size(partition_usage.free))+"\n"
    s7="Percentage: "+str(partition_usage.percent)+"%\n"
    with open("DISK_INFO.txt","a") as wh:
        print("Name: Filesystem and Size", file=wh)
        wh.write(s1)
        wh.write(s2)
        wh.write(s3)
        wh.write(s4)
        wh.write(s5)
        wh.write(s6)
        wh.write(s7)
    wh.close()
f = open("DISK_INFO.txt","r+")
for line in f:
    fields = line.split(":")
    f1=fields[0]
    f2=fields[1]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
f.truncate(0)
# add a page break to start a new page
#doc.add_page_break()
	
doc.add_heading(' Memory ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The following table shows the current status of memory utilization by Linux system and provides useful analytical information about memory")
doc_para.add_run("\n ")
svmem=psutil.virtual_memory()

### Table ###
table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Memory"
heading_cells[1].text = "Memory Information"

row = table.rows[1]
row.cells[0].text = "Total Memory"
row.cells[1].text = (f"Total:{get_size(svmem.total)}")

row = table.rows[2]
row.cells[0].text = "Available Memory"
row.cells[1].text = (f"Available:{get_size(svmem.available)}")

row = table.rows[3]
row.cells[0].text = "Used Memory"
row.cells[1].text = (f"Used:{get_size(svmem.used)}")

row = table.rows[4]
row.cells[0].text = "Percentage"
row.cells[1].text = (f"Percentage:{svmem.percent}%")


doc.add_heading(' SWAP Memory ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The following table shows the current status of virtual memory utilization by Linux system and provides useful analytical information about memory")
doc_para = doc.add_paragraph()
swap=psutil.swap_memory()

### Table ####
table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = "SWAP"
heading_cells[1].text = "Swap Memory Information"

row = table.rows[1]
row.cells[0].text = "Total Memory"
row.cells[1].text = (f"Total:{get_size(swap.total)}")

row = table.rows[2]
row.cells[0].text = "Available Memory"
row.cells[1].text = (f"Free:{get_size(swap.free)}")

row = table.rows[3]
row.cells[0].text = "Used Memory"
row.cells[1].text = (f"Used:{get_size(swap.used)}")

row = table.rows[4]
row.cells[0].text = "Percentage"
row.cells[1].text = (f"Percentage:{swap.percent}%")

#doc.add_page_break()
#doc_para.add_run("\n")
#doc_para = doc.add_paragraph(" ")

doc.add_heading(' Network ', 2)
doc_para = doc.add_paragraph()
#doc_para = doc.add_paragraph("This section provides information about IP configuration in the system.")
net_io = psutil.net_io_counters()
if_addrs = psutil.net_if_addrs()
for interface_name, interface_addresses in if_addrs.items():
    for address in interface_addresses:
        if str(address.family) == 'AddressFamily.AF_INET':
            s2="Interface;  "+str(interface_name)+"\n"
            s3="IP Address;  "+str(address.address)+"\n"
            s4="Netmask;  "+str(address.netmask)+"\n"
            s5="Broadcast IP;  "+str(address.broadcast)+"\n"
            with open("NETWORK_INFO.txt","a") as wh:
                print("Name;  Network Information; Network-IO-Byte-Receive; Network-IO-Byte-Sent", file=wh)
                wh.write(s2)
                wh.write(s3)
                wh.write(s4)
                wh.write(s5)
            wh.close()
        elif str(address.family) == 'AddressFamily.AF_PACKET':
            sx="MAC Address;  "+str(address.address)+"\n"
            with open("NETWORK_INFO.txt","a") as wh:
                wh.write(sx)
            wh.close()
f = open("NETWORK_INFO.txt","r+")
for line in f:
    fields = line.split(";")
    f1=fields[0]
    f2=fields[1]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
f.truncate(0)

#doc_para.add_run("\n")
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph()
doc.add_heading(' Network Route', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about Route in the system. ")
doc_para.add_run("\n ")
r=os.popen("/usr/sbin/route -n|grep -Ev 'Kernel'").read()
with open("ROUTE.txt","w") as wh:
        wh.write(r)
wh.close()
f = open("ROUTE.txt","r")
for line in f:
    fields = line.strip().split()
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    f4=fields[3]
    f5=fields[4]
    f6=fields[5]
    f7=fields[6]
    f8=fields[7]
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
    row.cells[3].text = (f4)
    row.cells[4].text = (f5)
    row.cells[5].text = (f6)
    row.cells[6].text = (f7)
    row.cells[7].text = (f8)
f.close()

doc.add_heading(' Services and Ports ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The following table is the true information about the ports used by application and it's associated services.")
doc_para.add_run("\n ")
se=os.popen("netstat -tulp|grep -Ev Active|grep LISTEN").read()
with open("service2.txt","w") as wh:
    print("Proto  Recv-Q  Send-Q  Local-Address  Foreign-Address  State  PID/Program name", file=wh)
    wh.write(se)
wh.close()

s2 = open("service2.txt","r")
for line in s2:
    fields = line.strip().split()
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    f4=fields[3]
    f5=fields[4]
    f6=fields[5]
    f7=fields[6]

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
    row.cells[3].text = (f4)
    row.cells[4].text = (f5)
    row.cells[5].text = (f6)
    row.cells[6].text = (f7)
s2.close()

doc.add_page_break()

doc.add_heading(' RHEL')
doc.add_heading(' Current Login Details', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section contains the Current login information in RHEL server ")
doc_para.add_run("\n ")
login=os.popen("/usr/bin/who|grep -Ev tty").read()
with open("LOGIN.txt","w") as wh:
    print("Login  Terminal  Login_Date  Login_Time  Remote_Server", file=wh)
    wh.write(login)
wh.close()
f = open("LOGIN.txt","r")
for line in f:
    login = line.split()
    f1=login[0]
    f2=login[1]
    f3=login[2]
    f4=login[3]
    f5=login[4]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
    row.cells[3].text = (f4)
    row.cells[4].text = (f5)
f.close()

doc.add_heading(' SELinux', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("SELinux should be enabled in Permissive mode on all systems as per the Atos Global standard. Current status can be found by running the following command:")
doc_para.add_run("\n ")
selinux=os.popen("sestatus").read()
with open("selinux.txt","w") as wh:
    wh.write(selinux)
wh.close()
f = open("selinux.txt","r")
for line in f:
    fields = line.strip().split(":")
    f2=fields[1]
    f3=f2.strip()
f.close()
table = doc.add_table(rows=2, cols=1)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = "SELinux status:"

row = table.rows[1]
row.cells[0].text = (f3)

#doc.add_page_break()

doc.add_heading(' Firewall', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The system should be configured to use Firewalld to allow and deny access.")
doc_para.add_run("\n ")
firewall=os.popen("firewall-cmd --state").read()
table = doc.add_table(rows=2, cols=1)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Firewalld status:"
row = table.rows[1]
row.cells[0].text = (firewall)

doc_para = doc.add_paragraph()

#doc_para.add_run("Current Firewalld Configuration").bold = True
fd=os.popen("firewall-cmd --zone=public --list-all").read()
table = doc.add_table(rows=2, cols=1)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Firewalld Configuration:"
row = table.rows[1]
row.cells[0].text = (fd)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Here services can be enabled to allow incoming connections. The above shows the default of primary interface firewalled and only port 22 open.")

doc.add_page_break()

doc.add_heading(' Target/Runlevel ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The default run level should be set to be multi-user.target")
doc_para = doc.add_paragraph()
doc_para.add_run("\n ")
doc_para.add_run("Verify default runlevel").bold = True
doc_para.add_run("\n ")
#r=os.popen("/bin/bash sed_task.sh").read()
#r=int(r)
#if(r<=6):
#    rn=os.popen("who -r").read()
#else:
rn=os.popen("systemctl get-default").read()
rn1="systemctl set-default multi-user.target"
#rn1=os.popen("systemctl set-default multi-user.target").read()
table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Default Runlevel"
heading_cells[1].text = "Command To Change The Runlevel"
row = table.rows[1]
row.cells[0].text = (rn)
row.cells[1].text = (rn1)

doc.add_page_break()

doc.add_heading(' Chrony/NTP', 2)
doc_para = doc.add_paragraph("chrony is an implementation of the Network Time Protocol. It's a replacement for the ntpd.")
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("The output displays information about the current time sources that chronyd is accessing.")
doc_para = doc.add_paragraph(".")
doc_para.add_run("\n ")
ntp=os.popen("chronyc tracking").read()
with open("chrony.txt","w") as wh:
    print("Name  :  Tracking", file=wh)
    wh.write(ntp)
wh.close()
ntp1 = open("chrony.txt","r")
for line in ntp1:
    fields = line.split(":")
    f1=fields[0]
    f2=fields[1]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
ntp1.close()

#doc.add_page_break()

doc.add_heading(' TimeZone', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("This section provides information about current TimeZone in the system. ")
doc_para.add_run("\n ")
timezone=os.popen("timedatectl|grep 'Time zone'").read()
with open("TIMEZONE.txt","w") as wh:
    wh.write(timezone)
wh.close()
f = open("TIMEZONE.txt","r")
for line in f:
    fields = line.strip().split(":")
    f1=fields[0]
    f2=fields[1]
#f.close()
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = True
# populate header row --------
    #heading_cells = table.rows[0].cells
    #heading_cells[0].text = "TimeZone"
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
f.close()

doc.add_page_break()

doc.add_heading(' Cluster ', 2)
#doc.add_heading(' Cluster Design', 3)
#doc.add_picture('/root/python/test_doc/Work_Script/Cluster.JPG',width=Inches(5.0))
#doc.add_heading (' Architecture Description', 4)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("RedHat Cluster are often used to support mission-critical services in the enterprise. The major goal of cluster is to keep services as available as possible by eliminating bottlenecks and single points of failure. ")
doc_para = doc.add_paragraph()
doc_para.add_run("\n ")
doc_para.add_run("Cluster Status:").bold = True
#r=os.popen("/bin/bash sed_task.sh").read()
#r=int(r)
#if(r<=6):
#    cluster=os.popen("clustat").read()
#else:

#import pathlib
#file = pathlib.Path("/etc/corosync/corosync.conf")
#if file.exists ():
#    ha=os.popen("pcs cluster status").read()
#    with open("pcs.txt","w") as wh:
#        wh.write(ha)
#    wh.close()
cluster = open("pcs.txt","r")
for line in cluster:
    fields = line.split()
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (fields)
#else:
#    doc_para.add_run("Cluster is not Configured").bold = True

doc.add_heading(' Log File ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Log files are files that contain messages about the system, including the kernel, services, and applications running on it. There are different log files for different information. For example:")
doc_para = doc.add_paragraph("/var/log/messages: This file has all the global system messages located inside, including the messages that are logged during system startup. Depending on how the syslog config file is sent up, there are several things that are logged in this file including mail, cron, daemon, kern, auth, etc.", style='List Bullet')
doc_para = doc.add_paragraph("/var/log/secure: Contains information related to authentication and authorization privileges. For example, sshd logs all the messages here, including unsuccessful login", style='List Bullet')
doc_para = doc.add_paragraph("/var/log/boot.log: Contains information that are logged when the system boots", style='List Bullet')
doc_para = doc.add_paragraph()
m=os.popen('ls -lh /var/log/messages|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Live Log File"}\'').read()
s=os.popen('ls -lh /var/log/secure|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Live Log File"}\'').read()
b=os.popen('ls -lh /var/log/boot.log|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Live Log File"}\'').read()
m1=os.popen('ls -lh /var/log/messages-*|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Archive Log File"}\'').read()
s1=os.popen('ls -lh /var/log/secure-*|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Archive Log File"}\'').read()
b1=os.popen('ls -lh /var/log/boot.log-*|awk \'{print $9 ";" $6,$7,$8 ";" $5 ";" echo "Archive Log File"}\'').read()
with open("LOG.txt","w") as wh:
    print("Log_File  ;  Date  ;  Size  ;  File Type",file=wh)
    wh.write(m+s+b+m1+s1+b1)
wh.close()
f = open("LOG.txt","r")
for line in f:
    fields = line.split(";")
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    f4=fields[3]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
    row.cells[3].text = (f4)

doc.add_page_break()

doc.add_heading(' Atos Technology Framework ')
doc_para = doc.add_paragraph("The Atos Technology Framework provides a tooling solution for the Atos Service Management Model (ASMM) and the associated processes and consistently manages the interactions between all components, the Services and all users based on a flexible IT architecture. Henkel tooling will be removed, and Atos tooling will be installed.")
doc.add_heading(' Monitoring Tool ', 2)
doc_para = doc.add_paragraph()
#doc_para.add_run("Verify CMF Nagios and ASE Agent Status").bold = True
doc_para = doc.add_paragraph()
nacl = os.popen("rpm -q atos-cmf-client-nacl").read()
if len(nacl) != 0:
    #nacl_service=os.popen("su - nagios -c NaCl/NaCl").read()
    nacl_service=os.popen("cat /home/nagios/NaCl/NaCl1.log|grep 'WARNING'|awk -F: '{print $2}'").read()
    ase = os.popen("rpm -q ase").read()
    ase_service=os.popen("systemctl is-active ase").read()
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.allow_autofit = True
    # populate header row -------
    heading_cells = table.rows[0].cells
    heading_cells[0].text = "CMF Nagios Agent"
    heading_cells[1].text = "CMF Nagios Agent Status"
    heading_cells[2].text = "ASE Agent"
    heading_cells[3].text = "ASE Agent Status"
    row = table.rows[1]
    row.cells[0].text = nacl
    row.cells[1].text = nacl_service
    row.cells[2].text = ase
    row.cells[3].text = ase_service
else:
    doc_para.add_run("atos-cmf-client-nacl package is not installed").bold = True

doc.add_heading(' Patching Tool ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Linux servers will be patched using BMC blade logic tool. BMC BladeLogic agent RPM will be installed in the server and GDTS BladeLogic team will add that server into the management. RPM should be received from GDTS team")
doc_para.add_run("\n")
bladelogic=os.popen("rpm -qa |grep BladeLogic").read()
if len(bladelogic) != 0:
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.allow_autofit = True
    # populate header row -------
    heading_cells = table.rows[0].cells
    heading_cells[0].text = "BMC Agent"
    row = table.rows[1]
    row.cells[0].text = (bladelogic)
    doc_para = doc.add_paragraph()
    doc_para.add_run("\n ")
    doc_para.add_run("BMC Agent Status").bold = True
    doc_para.add_run("\n")
    bladelogic1=os.popen("ps aux|grep rscd |grep -v grep").read()
    with open("bladelogic.txt","w") as wh:
        print("USER       PID %CPU %MEM    VSZ   RSS TTY      STAT START   TIME COMMAND", file=wh)
        wh.write(bladelogic1)
    wh.close()
    b1 = open("bladelogic.txt","r")
    for line in b1:
        fields = line.split()
        f1=fields[0]
        f2=fields[7]
        f3=fields[10]
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.allow_autofit = True
        row = table.rows[0]
        row.cells[0].text = (f1)
        row.cells[1].text = (f2)
        row.cells[2].text = (f3)
    b1.close()
else:
    doc_para.add_run("BladeLogic package is not installed").bold = True
#doc.add_page_break()

doc.add_heading(' Backup Tool ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("EMC Networker agent will be installed on Linux servers")
doc_para.add_run("\n")
networker=os.popen("rpm -qa |grep -i lgto").read()
if len(networker) != 0:
    nwservice=os.popen("systemctl is-active networker").read()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.allow_autofit = True
    # populate header row -------
    heading_cells = table.rows[0].cells
    heading_cells[0].text = "NetWorker Client"
    heading_cells[1].text = "NetWorker Service"
    row = table.rows[1]
    row.cells[0].text = (networker)
    row.cells[1].text = (nwservice)
else:
    doc_para.add_run("Networker package is not installed").bold = True

doc.add_heading(' System Recovery Tool ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Relax-and-Recover(ReaR) is a recovery and system migration utility. The utility produces a bootable image and restores from backup using this image. It also allows to restore to different hardware and can therefore be used as a migration utility as well.")
doc_para.add_run("\n")
rear=os.popen("/usr/sbin/rear -V").read()
table = doc.add_table(rows=2, cols=1)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].text = "REAR Version"
row = table.rows[1]
row.cells[0].text = (rear)

doc.add_page_break()

doc.add_heading(' Health Checkup ')
doc_para = doc.add_paragraph("The following data table is very important during hardening and onboarding of RHEL Server instance. It provides Precise information about server health based on Information Security Standards. ")
doc_para.add_run("\n ")
os.system("/usr/bin/python3 Linux_Helath_Check.py")
ht = open("HEALTH.txt","r")
for line in ht:
    fields = line.split()
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
ht.close()

doc.add_page_break()

doc.add_heading(' Atos TSS ')
doc_para = doc.add_paragraph("Atos standard build procedure will be used to harden the system. Every new server installed should be secured following the Unix Security Standards.  Note that there may be customer requirements which mean that configurations listed in the policy may need to be altered.  Any such alterations should be documented as exceptions to the Unix Security Standards in the server documentation.")
doc_para.add_run("\n ")
os.system("/usr/bin/python3 tss-script-new.py")
ts = open("TSS.txt","r")
for line in ts:
    #fields = line.strip(" ").split(" ")
    fields = line.split(":")
    #print(fields)
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    f4=fields[3]
    f5=fields[4]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
    row.cells[3].text = (f4)
    row.cells[4].text = (f5)
ts.close()

doc.add_page_break()

doc.add_heading(' Risk Analysis ')
doc_para = doc.add_paragraph("The following data table is very important during hardening and onboarding of RHEL Servers. It provides Deep Dive Risk information that must be taken into consideration . ")
doc_para.add_run("\n ")
os.system("/usr/bin/python3 risk-ana-script-new.py")
rk = open("RISK_ANA.txt","r")
for line in rk:
    #fields = line.strip(" ").split(" ")
    fields = line.split(":")
    #print(fields)
    f1=fields[0]
    f2=fields[1]
    f3=fields[2]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.allow_autofit = True
    row = table.rows[0]
    row.cells[0].text = (f1)
    row.cells[1].text = (f2)
    row.cells[2].text = (f3)
rk.close()

doc.add_page_break()

doc.add_heading(' Backup ')
doc_para = doc.add_paragraph("EMC Networker agent will be installed on servers, under '/nsr' to take all drives/partition backup and it will be backed up in accordance with the existing Atos standard regime.")
doc_para = doc.add_paragraph("RMAN backup tool will be used to take Oracle DB backups.", style='List Bullet')
doc_para = doc.add_paragraph("A full backup will be done once a week, incremental backups will be done for the rest of the week. Backup retention period will be one month or as per Atos backup policy standard.", style='List Bullet')
doc_para.add_run("\n ")
mount=os.popen("df -h|grep -v 'File'|awk '{print $6}'|grep -Ev 'run|dev|sys|mnt'").read()
vg=os.popen("vgs").read()
with open("vg.txt","w") as wh:
    wh.write(vg)
    wh.close()
v = open("vg.txt","r")
for line in v:
    fields = line.split()
    f1=fields[5]
    f2 = ""+str(f1)+""
    f3 = ""+str(f2.replace('<', ' ').replace('VSize',' '))+""
#    print(f3)
v.close()
table = doc.add_table(rows=3, cols=5)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Service Component"
heading_cells[1].text = "Devices"
heading_cells[2].text = "Data Volumes (GB)"
heading_cells[3].text = "Agent Type"
heading_cells[4].text = "Archive Frequency"
row = table.rows[1]
row.cells[0].text = "Operating system , all drives/partition Local"
row.cells[1].text = (mount)
row.cells[2].text = (f3)
row.cells[3].text = "EMC Networker"
row.cells[4].text = "Daily Incremental\Weekly Full"
row = table.rows[2]
row.cells[0].text = "Operating system , all drives/partition on SAN"
row.cells[1].text = "Not Applicable"
row.cells[2].text = "Not Applicable"
row.cells[3].text = "EMC Networker"
row.cells[4].text = "Daily Incremental\Weekly Full"

doc.add_page_break()

doc.add_heading(' System Recovery ')
doc_para = doc.add_paragraph("System recovery processes are used to restore a server after some failure be it hardware or software, it can be regarded as part of Disaster Recovery (DR) or as a separate entity if a customer does not have DR. There are several methods of system recovery listed below:")
doc_para = doc.add_paragraph()
doc_para.add_run("\n")
doc_para.add_run("Relax and Recover (REAR)").bold = True
doc_para.add_run("\n")
doc_para.add_run("Relax-and-Recover produces a bootable image which can recreate the system’s original storage layout. Once that is done it initiates a restore from backup. Since the storage layout can be modified prior to recovery, and dissimilar hardware and virtualization is supported, Relax-and-Recover offers the flexibility to be used for complex system migrations.")
doc_para = doc.add_paragraph()
doc_para.add_run("\n")
doc_para.add_run("Enterprise system restore (NetBackup / Networker etc) ").bold = True
doc_para.add_run("\n")
doc_para.add_run("An enterprise solution can be used to restore a system such as NetBackup (Symantec) or Networker (EMC).")

doc.add_page_break()

doc.add_heading(' Redhat Licensing ')
doc_para = doc.add_paragraph("Red Hat Subscriptions (or licensing) are available in many forms, the 2 common ones are:")
doc_para = doc.add_paragraph("RED HAT ENTERPRISE LINUX SERVER", style='List Bullet')
doc_para = doc.add_paragraph("RED HAT ENTERPRISE LINUX FOR VIRTUAL DATACENTERS", style='List Bullet')
doc_para = doc.add_paragraph()
doc_para.add_run("\n")
doc_para.add_run("Table 1. Service-level agreements for Red Hat Enterprise Linux subscriptions.").bold = True
doc_para.add_run("\n")
table = doc.add_table(rows=8, cols=4)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Service"
heading_cells[1].text = "Self-support"
heading_cells[2].text = "Standard"
heading_cells[3].text = "Premium"
row = table.rows[1]
row.cells[0].text = "Hours of coverage"
row.cells[1].text = "None"
row.cells[2].text = "Standard business"
row.cells[3].text = "Standard business hours (24x7 for Severity 1 and Severity 2)"
row = table.rows[2]
row.cells[0].text = "Support channel"
row.cells[1].text = "None"
row.cells[2].text = "Web and phone"
row.cells[3].text = "NA"
row = table.rows[3]
row.cells[0].text = "Number of cases"
row.cells[1].text = "None"
row.cells[2].text = "Unlimited"
row.cells[3].text = "Unlimited"
row = table.rows[4]
row.cells[0].text = "Severity 1 (Urgent)"
row.cells[1].text = "1 business hour"
row.cells[2].text = "1 hour"
row.cells[3].text = "1 hour"
row = table.rows[5]
row.cells[0].text = "Severity 2 (High)"
row.cells[1].text = "4 business hours"
row.cells[2].text = "2 business hours"
row.cells[3].text = "4 hours"
row = table.rows[6]
row.cells[0].text = "Severity 3 (Medium)"
row.cells[1].text = "1 business day"
row.cells[2].text = "4 business hours"
row.cells[3].text = "8 business hours"
row = table.rows[7]
row.cells[0].text = "Severity 4 (Low)"
row.cells[1].text = "2 business days"
row.cells[2].text = "8 business hours"
row.cells[3].text = "2 business days or as agreed"

doc.add_page_break()

doc.add_heading(' Redhat LifeCycle ')
doc.add_heading(' End of Life ', 2)
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph("Red Hat offers subscription services for each major release of Red Hat Enterprise Linux throughout four life-cycle phases—called Full Support, Maintenance Support 1, Maintenance Support 2, and an Extended Life Phase.")
doc_para.add_run("\n")
doc_para = doc.add_paragraph(" Please visit RedHat link mentioned below for Red Hat Enterprise Linux Life Cycle policies.")
#doc_para.add_run("\n")
p = doc.add_paragraph()
add_hyperlink(p, 'Red Hat Enterprise Linux Life Cycle', "https://access.redhat.com/support/policy/updates/errata")
doc_para.add_run("\n")
doc_para = doc.add_paragraph()
doc_para = doc.add_paragraph(" The content has been taken from offline")
doc_para = doc.add_paragraph()
table = doc.add_table(rows=2, cols=4)
table.style = "Table Grid"
table.allow_autofit = True
# populate header row -------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Products Released"
heading_cells[1].text = "Lifecycle Start Date"
heading_cells[2].text = "Full support ends"
heading_cells[3].text = "Maintenance Support ends"

row = table.rows[1]
row.cells[0].text = "Redhat Enterprise Linux Server 8"
row.cells[1].text = "May 7, 2019"
row.cells[2].text = "May 31, 2024"
row.cells[3].text = "May 31, 2029"

doc.add_page_break()
e = datetime.datetime.now()
d=e.strftime("%Y-%m-%d_%H-%M-%S")
doc.save("Redhat_LLD_"+str(d)+".docx")
